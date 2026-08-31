"""P1-1 weekly summary report: fetch via MCP, build a Word document.

The report is *generated material*, not a chat transcript (plan 5.4):
deterministic fetch of aggregated calibers from the取数 service, laid out
as a leader-style summary, exported as .docx via python-docx (the
haitun-workspace write_word pattern, with w:eastAsia set so Chinese renders
evenly).

The BFF talks to the MCP service directly with its own token — the report
never depends on a model turn, so it is stable and cheap to regenerate.
"""

from __future__ import annotations

import io
import json
from contextlib import suppress
from dataclasses import dataclass, field
from datetime import date
from typing import Any

import httpx
from docx import Document
from docx.shared import Pt
from mcp import ClientSession
from mcp.client.streamable_http import streamable_http_client

from .config import BffConfig


@dataclass
class WeeklySummaryData:
    snapshot_note: str = ""
    caliber: str = ""
    status_rows: list[dict[str, Any]] = field(default_factory=list)
    board_rows: list[dict[str, Any]] = field(default_factory=list)
    freshness_rows: list[dict[str, Any]] = field(default_factory=list)
    stale_rows: list[dict[str, Any]] = field(default_factory=list)
    group_rows: list[dict[str, Any]] = field(default_factory=list)
    group_stale_rows: list[dict[str, Any]] = field(default_factory=list)
    never_reported_rows: list[dict[str, Any]] = field(default_factory=list)
    milestone_row: dict[str, Any] = field(default_factory=dict)
    year_goal_rows: list[dict[str, Any]] = field(default_factory=list)


async def fetch_weekly_summary(config: BffConfig) -> WeeklySummaryData:
    """Ask the取数 service for the aggregates the report is built from.

    Every call is guarded: one failing caliber must not sink the whole
    report — the section just stays empty instead.
    """
    data = WeeklySummaryData()
    timeout = httpx.Timeout(30.0)
    headers = {
        "Authorization": f"Bearer {config.mcp_token}",
        "Accept": "application/json, text/event-stream",
    }
    if not config.mcp_token:
        # An empty token must not produce an "Authorization: Bearer " header —
        # that is an illegal header value (httpx raises LocalProtocolError).
        headers.pop("Authorization", None)
    # trust_env=False: the取数 service is loopback/internal — a machine-level
    # proxy must never hijack it.
    async with (
        httpx.AsyncClient(headers=headers, timeout=timeout, trust_env=False) as http_client,
        streamable_http_client(config.mcp_url, http_client=http_client) as parts,
        ClientSession(parts[0], parts[1]) as session,
    ):
        await session.initialize()

        async def call(name: str, arguments: dict[str, Any]) -> dict[str, Any]:
            result = await session.call_tool(name, arguments)
            # SDK field naming drifted between versions (structuredContent vs
            # structured_content) — accept both.
            structured = getattr(result, "structured_content", None) or getattr(result, "structuredContent", None)
            if isinstance(structured, dict) and set(structured) == {"result"}:
                raw = structured["result"]
                if isinstance(raw, str):
                    try:
                        return json.loads(raw)
                    except json.JSONDecodeError:
                        return {"ok": False}
            return structured or {}

        async def safe(name: str, arguments: dict[str, Any]) -> dict[str, Any]:
            try:
                return await call(name, arguments)
            except Exception:
                return {"ok": False}

        status = await safe("weekly_aggregate", {"group_by": "status"})
        if status.get("ok"):
            data.status_rows = _rows(status)
            data.caliber = str(status.get("caliber", ""))
            data.snapshot_note = str(status.get("snapshot_note", ""))

        board = await safe("weekly_aggregate", {"group_by": "board"})
        if board.get("ok"):
            data.board_rows = _rows(board)

        freshness = await safe("weekly_freshness", {})
        if freshness.get("ok"):
            data.freshness_rows = _rows(freshness)

        # 进展陈旧度分档 (默认档: 30/90/180 天/从未, 5 行计数表).
        stale = await safe("weekly_freshness_distribution", {})
        if stale.get("ok"):
            data.stale_rows = _rows(stale)

        # 专项组完成率对比, 最低在前 (caliber 推荐 order_by=finish_rate + ascending).
        groups = await safe(
            "weekly_aggregate",
            {"group_by": "project_group", "order_by": "finish_rate", "ascending": True},
        )
        if groups.get("ok"):
            data.group_rows = _rows(groups)

        # 组级滞后占比 (stale_days + by=project_group).
        group_stale = await safe("weekly_freshness_distribution", {"stale_days": 30, "by": "project_group"})
        if group_stale.get("ok"):
            data.group_stale_rows = _rows(group_stale)

        # 从未上报进展的任务清单 (前 10 条 + 总数由 caller 说明).
        never = await safe("weekly_progress_coverage", {"scope": "never_reported"})
        if never.get("ok"):
            data.never_reported_rows = _rows(never)

        # 里程碑完成率 — 单行汇总, 失败则整节留空.
        milestone = await safe("weekly_milestone_stats", {})
        if milestone.get("ok"):
            milestone_rows = _rows(milestone)
            if milestone_rows:
                data.milestone_row = milestone_rows[0]

        # 年度目标按建单年份分布 (默认档).
        year_goal = await safe("weekly_year_goal_stats", {})
        if year_goal.get("ok"):
            data.year_goal_rows = _rows(year_goal)

    return data


# The OOXML eastAsia font attribute name — kept as a constant so the long
# namespace URL does not break the line-length rule everywhere it is set.
_W_NS_EAST_ASIA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"


def _rows(payload: dict[str, Any]) -> list[dict[str, Any]]:
    rows = payload.get("rows")
    if not isinstance(rows, list):
        return []
    return [row for row in rows if isinstance(row, dict)]


def _row_value(row: dict[str, Any], *keys: str) -> Any:
    """First non-empty value among the candidate keys (caliber field names
    vary slightly across scopes)."""
    for key in keys:
        value = row.get(key)
        if value is not None and str(value).strip() != "":
            return value
    return ""


def _status_count(rows: list[dict[str, Any]], label_fragment: str) -> int:
    """Count from a status table row whose label contains the fragment."""
    for row in rows:
        label = str(_row_value(row, "status_label", "group_name", "label"))
        if label_fragment in label:
            try:
                return int(_row_value(row, "cnt", "count"))
            except TypeError, ValueError:
                return 0
    return 0


def _first_group_metric(rows: list[dict[str, Any]], metric_key: str, *, lowest: bool) -> tuple[str, str]:
    """(group name, metric) of the lowest/highest row — scans numerically:
    row order is not guaranteed to match the metric's order."""
    best: tuple[float, str, str] | None = None
    for row in rows:
        raw = str(_row_value(row, metric_key)).replace("%", "").strip()
        try:
            value = float(raw)
        except ValueError:
            continue
        if best is None or (lowest and value < best[0]) or (not lowest and value > best[0]):
            name = str(_row_value(row, "bucket", "group_name", "project_group", "board_name"))
            best = (value, name, raw)
    if best is None:
        return "", ""
    return best[1], best[2]


def build_summary_document(data: WeeklySummaryData) -> bytes:
    """Lay the fetched data out as a leader-style weekly summary .docx:
    总体概况 / 进展与时效 / 专项组完成情况 / 风险与滞后 / 关注建议."""
    doc = Document()
    _set_east_asia(doc)

    title = doc.add_heading("", level=0)
    _run_with_font(title, "周报总结", bold=True, size=Pt(26))

    meta = doc.add_paragraph()
    _run_with_font(meta, f"生成日期:{date.today().isoformat()}", bold=True)
    meta_second = doc.add_paragraph()
    _run_with_font(meta_second, data.snapshot_note or "数据快照见正文", size=Pt(9))
    if data.caliber:
        para = doc.add_paragraph()
        _run_with_font(para, f"口径:{data.caliber}", italic=True, size=Pt(9))

    total = sum(_status_count(data.status_rows, label) for label in ("未开始", "进行中", "已完成", "已停用"))
    finished = _status_count(data.status_rows, "已完成")
    in_flight = _status_count(data.status_rows, "进行中")

    # ── 一、总体概况 ──
    doc.add_heading("一、总体概况", level=1)
    if total:
        summary = (
            f"截至数据快照,正式任务共 {total} 条:已完成 {finished} 条({finished / total:.1%})、"
            f"进行中 {in_flight} 条,其余为未开始与已停用。"
        )
        _run_with_font(doc.add_paragraph(), summary)
    if data.status_rows:
        doc.add_paragraph("状态分布:")
        _style_table(_add_table(doc, data.status_rows))
    if data.board_rows:
        doc.add_paragraph("看板分布:")
        _style_table(_add_table(doc, data.board_rows))

    # ── 二、进展与时效 ──
    doc.add_heading("二、进展与时效", level=1)
    if data.freshness_rows:
        doc.add_paragraph("各看板最新进展时间:")
        _style_table(_add_table(doc, data.freshness_rows))
    if data.stale_rows:
        doc.add_paragraph("进展陈旧度分档(全库):")
        _style_table(_add_table(doc, data.stale_rows))

    # ── 三、专项组完成情况 ──
    doc.add_heading("三、专项组完成情况", level=1)
    if data.group_rows:
        lowest_name, lowest_rate = _first_group_metric(data.group_rows, "finish_rate_pct", lowest=True)
        highest_name, highest_rate = _first_group_metric(data.group_rows, "finish_rate_pct", lowest=False)
        if lowest_name and lowest_rate:
            sentence = f"完成率最低的是{lowest_name}({lowest_rate}%),最高的是{highest_name}({highest_rate}%)。"
            _run_with_font(doc.add_paragraph(), sentence)
        doc.add_paragraph("各专项组完成率对比(按完成率升序):")
        _style_table(_add_table(doc, data.group_rows))
    if data.milestone_row:
        finished_ms = _row_value(data.milestone_row, "finished", "finished_count")
        total_ms = _row_value(data.milestone_row, "total", "total_count")
        rate_ms = _row_value(data.milestone_row, "finish_rate_pct", "finish_rate")
        _run_with_font(
            doc.add_paragraph(),
            f"里程碑:{total_ms} 个,已完成 {finished_ms} 个(完成率 {rate_ms}%)。",
        )
    if data.year_goal_rows:
        doc.add_paragraph("年度目标按建单年份分布:")
        _style_table(_add_table(doc, data.year_goal_rows))

    # ── 四、风险与滞后 ──
    doc.add_heading("四、风险与滞后", level=1)
    if data.group_stale_rows:
        stale_name, stale_pct = _first_group_metric(data.group_stale_rows, "stale_pct", lowest=False)
        if stale_name and stale_pct:
            _run_with_font(
                doc.add_paragraph(),
                f"滞后占比最高的是{stale_name}({stale_pct}%),为各组之最,需重点治理。",
            )
        doc.add_paragraph("各组滞后与活跃情况:")
        _style_table(_add_table(doc, data.group_stale_rows))
    if data.never_reported_rows:
        shown = data.never_reported_rows[:10]
        total_nr = len(data.never_reported_rows)
        group_only = sum(1 for row in data.never_reported_rows if row.get("has_group_history"))
        truly_none = total_nr - group_only
        doc.add_paragraph(
            f"从未上报进展的任务(按 task_progress 有无已发布行判定,共 {total_nr} 条;"
            f"其中集团看板 {group_only} 条的成效记于集团历史表,不算漏报;"
            f"两表均无的 {truly_none} 条。列前 10):"
        )
        _style_table(_add_table(doc, shown))

    # ── 五、关注建议 ──
    doc.add_heading("五、关注建议", level=1)
    suggestions: list[str] = []
    if total and finished:
        suggestions.append(f"1. 总体完成率 {finished / total:.1%},过半任务仍在推进,建议按周排期滚动验收。")
    if data.group_stale_rows:
        stale_name, stale_pct = _first_group_metric(data.group_stale_rows, "stale_pct", lowest=False)
        if stale_name:
            suggestions.append(f"2. 滞后集中在{stale_name}等组(最高 {stale_pct}%),建议专项督办并限期补报。")
    if data.never_reported_rows:
        suggestions.append(f"3. {len(data.never_reported_rows)} 条任务从未上报进展,需逐条确认口径与填报责任。")
    suggestions.append("4. 核对下一批进展提交单是否按期落库,避免对外期号与内部进展脱节。")
    for suggestion in suggestions:
        _run_with_font(doc.add_paragraph(), suggestion)

    doc.add_paragraph()
    note = doc.add_paragraph()
    _run_with_font(note, "数据来源:演示库(weekly_mock),非集团真实周报。", bold=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_east_asia(doc: Document) -> None:
    """w:eastAsia on every used style keeps Chinese fonts even across viewers.

    Without it, heading styles fall back to MS Gothic for CJK and the
    document comes out in mismatched typefaces. Runs get the font set
    explicitly too, so numbers and Latin do not slip back to Calibri.
    """
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Microsoft YaHei"
        if style._element.rPr is not None and style._element.rPr.rFonts is not None:
            style._element.rPr.rFonts.set(_W_NS_EAST_ASIA, "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)


def _run_with_font(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: Pt | None = None,
) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = size
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(_W_NS_EAST_ASIA, "Microsoft YaHei")


def _add_table(doc: Document, rows: list[dict[str, Any]]):
    if not rows:
        return doc.add_table(rows=0, cols=0)
    columns = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(columns))
    for index, name in enumerate(columns):
        table.rows[0].cells[index].text = str(name)
    for row in rows:
        cells = table.add_row().cells
        for index, name in enumerate(columns):
            cells[index].text = str(row.get(name, ""))
    return table


def _style_table(table: Any) -> None:
    with suppress(KeyError):
        table.style = "Light Grid Accent 1"


def build_summary_document_text(data: WeeklySummaryData) -> str:
    """Plain-text fallback preview (used by tests and debugging)."""
    lines = ["周报总结", f"快照:{data.snapshot_note}", f"口径:{data.caliber}"]
    for title, rows in (
        ("状态分布", data.status_rows),
        ("看板", data.board_rows),
        ("时效", data.freshness_rows),
        ("滞后分档", data.stale_rows),
        ("专项组完成率", data.group_rows),
        ("组级滞后", data.group_stale_rows),
        ("从未上报", data.never_reported_rows),
    ):
        lines.append(f"[{title}] {len(rows)} 行")
    if data.milestone_row:
        lines.append(f"[里程碑] {data.milestone_row}")
    if data.year_goal_rows:
        lines.append(f"[年度目标] {len(data.year_goal_rows)} 行")
    lines.append("数据来源:演示库(weekly_mock),非集团真实周报。")
    return "\n".join(lines)
